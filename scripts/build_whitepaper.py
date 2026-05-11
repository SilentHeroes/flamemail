#!/usr/bin/env python3
"""Build the Flamemail whitepaper as a docx, styled to match the leaflet_project format."""

from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/nicholas.kuntz/flamemail/Flamemail-Whitepaper.docx"
ARIAL = "Arial"
GRAY = RGBColor(0x55, 0x55, 0x55)


def style_run(run, *, bold=False, size=None, color=None, font=ARIAL):
    run.font.name = font
    if size is not None:
        run.font.size = size
    if bold:
        run.font.bold = True
    if color is not None:
        run.font.color.rgb = color


def add_centered(doc, text, *, bold=False, color=None, size=None, space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    style_run(run, bold=bold, color=color, size=size)
    return p


def add_section_heading(doc, text):
    """Numbered/named section heading: bold, with space before."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    style_run(run, bold=True)
    return p


def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    style_run(run, bold=True)
    return p


def add_body(doc, segments, *, space_after=Pt(0)):
    """segments: list of (text, bold) tuples — single body paragraph with mixed bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    if isinstance(segments, str):
        segments = [(segments, False)]
    for text, bold in segments:
        run = p.add_run(text)
        style_run(run, bold=bold)
    return p


def add_bullet(doc, segments):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    for text, bold in segments:
        run = p.add_run(text)
        style_run(run, bold=bold)
    return p


def main():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = ARIAL
    normal.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Emu(548640)
        section.bottom_margin = Emu(548640)
        section.left_margin = Emu(640080)
        section.right_margin = Emu(640080)

    # ---- Title block (centered) ----
    add_centered(
        doc,
        "Disposable Inboxes and Covert Two-Way Email Relay on Cloudflare's Edge",
        bold=True,
    )
    add_centered(
        doc,
        "A Cloudflare Workers architecture for ephemeral mail and unlinkable two-domain relay",
    )
    add_centered(
        doc,
        "Cloudflare · Technical White Paper",
        color=GRAY,
        space_after=Pt(2),
    )

    # ---- Overview ----
    add_section_heading(doc, "Overview")
    add_body(
        doc,
        "Email remains the universal contact channel, but the address itself is a durable identifier "
        "that follows users across services, leaks correspondence patterns to ISPs, and outlives the "
        "purpose it was created for. Flamemail is a Cloudflare Workers application that provides two "
        "complementary capabilities from one deployable unit: short-lived disposable inboxes that "
        "auto-destruct on schedule, and a covert two-way relay that lets two parties exchange email "
        "through a single shared inbox accessed from two unrelated domains. There are no servers to "
        "operate, no mail transport agents to harden, and no third-party SaaS in the data path. "
        "Inbound MX, outbound delivery, storage, real-time push, bot mitigation, and the web UI all "
        "run inside one Worker on Cloudflare's global network.",
    )

    # ---- 1 ----
    add_section_heading(doc, "1. One Worker, Three Handlers")
    add_body(
        doc,
        [
            ("The entire system is a single Worker exposing three handlers: ", False),
            ("fetch", True),
            (" for HTTP and WebSocket, ", False),
            ("email", True),
            (" for inbound MX delivery from Cloudflare Email Routing, and ", False),
            ("scheduled", True),
            (" for hourly cleanup. The Worker itself is stateless. State is partitioned across "
             "managed Cloudflare primitives so any data-center can serve any request without "
             "session affinity.",
             False),
        ],
    )
    add_bullet(doc, [("D1 (SQLite). ", True), ("Inbox metadata, email records, attachments, and the relay alias table.", False)])
    add_bullet(doc, [("R2. ", True), ("Raw .eml payloads, parsed body JSON, and individual attachment blobs keyed by inbox ID.", False)])
    add_bullet(doc, [("KV. ", True), ("Bearer session tokens and one-time WebSocket tickets with TTL expiry matching the inbox lifetime.", False)])
    add_bullet(doc, [("Durable Objects. ", True), ("Per-inbox hibernation WebSockets that fan out new-message events to connected clients in real time, with near-zero idle cost.", False)])
    add_bullet(doc, [("Email Routing + Email Sending. ", True), ("Catch-all MX intake on each registered domain and the send_email binding for outbound delivery, both native Cloudflare bindings.", False)])
    add_bullet(doc, [("Turnstile. ", True), ("Bot mitigation on inbox creation, relay creation, and admin login. Fails closed when not configured.", False)])
    add_bullet(doc, [("Static Assets. ", True), ("React SPA (Vite + Tailwind) served from the same Worker. No separate origin, no CORS surface.", False)])

    # ---- 2 ----
    add_section_heading(doc, "2. Inbound Email Pipeline")
    add_body(
        doc,
        "Mail addressed to any registered domain hits the Worker's email() handler via a catch-all "
        "Email Routing rule. The Worker normalizes the recipient (stripping plus-addressing), looks "
        "up the inbox or resolves it through the relay alias table, and enforces caps (10 MB per "
        "message, 100 messages per inbox, 10 attachments). postal-mime parses the MIME tree, "
        "metadata is batch-inserted into D1, raw and parsed payloads are written to R2, and the "
        "inbox's Durable Object broadcasts a notification to every connected WebSocket client. If "
        "the inbox is a relay with a registered notification address, an alert is dispatched via "
        "the send_email binding so neither party has to keep the tab open.",
    )

    # ---- 3 Relay ----
    add_section_heading(doc, "3. Relay Mode")
    add_body(
        doc,
        "Relay mode is the differentiating capability. Two parties, each on a different domain, "
        "exchange email through one shared inbox without their respective ISPs ever observing "
        "a message that traverses both networks. The two endpoint addresses look unrelated to any "
        "external observer.",
    )

    add_subsection_heading(doc, "3.1 Deterministic Derivation")
    add_body(
        doc,
        [
            ("Both parties enter the same passphrase independently. PBKDF2-SHA256 with 100,000 "
             "iterations derives a deterministic 8-character local part per domain, so Party A on ",
             False),
            ("mail.example.com", True),
            (" sees ", False),
            ("xk7f9m2q@example.com", True),
            (" while Party B on ", False),
            ("relay.example.net", True),
            (" sees ", False),
            ("yq3n8c1p@example.net", True),
            (". The passphrase is never persisted; only SHA-256(passphrase) is stored to make "
             "creation idempotent across both endpoints.",
             False),
        ],
    )

    add_subsection_heading(doc, "3.2 Routing and Aliasing")
    add_body(
        doc,
        [
            ("One primary inbox is created in D1. The secondary local part is registered in the ",
             False),
            ("relay_pairs", True),
            (" table as an alias. The email handler resolves alias → primary at delivery time, "
             "so both addresses land in the same inbox and both parties see the same thread. Either "
             "party can compose outbound mail from either address; replies route correctly without "
             "either side learning the other's domain.",
             False),
        ],
    )

    add_subsection_heading(doc, "3.3 Observability Surface")
    add_body(
        doc,
        "Party A's ISP only sees traffic to and from example.com. Party B's ISP only sees traffic "
        "to and from example.net. There is no network-level evidence that the two users are "
        "corresponding. Cascade deletion ties relay aliases to the primary inbox's expiry, so the "
        "linkage disappears with the data.",
    )

    # ---- 4 Use cases ----
    add_section_heading(doc, "4. Use Cases")
    add_bullet(doc, [
        ("Sign-up hygiene and SaaS evaluation. ", True),
        ("Issue a disposable inbox for a vendor trial, marketing list, or one-time verification. "
         "It expires on schedule and the address cannot be re-used to profile the recipient.",
         False),
    ])
    add_bullet(doc, [
        ("Journalist and source contact. ", True),
        ("A reporter publishes a relay address on one domain; a source contacts a different "
         "domain with the agreed passphrase. Neither side's mail provider can correlate them, "
         "and the channel self-destructs.",
         False),
    ])
    add_bullet(doc, [
        ("Incident response and red-team coordination. ", True),
        ("Stand up an out-of-band channel that does not touch corporate mail. Rotate the "
         "passphrase to rotate the channel; cleanup is automatic.",
         False),
    ])
    add_bullet(doc, [
        ("Vendor and M&A due diligence. ", True),
        ("Exchange sensitive attachments through a shared inbox that is not tied to either "
         "organization's primary domain or retention policy.",
         False),
    ])
    add_bullet(doc, [
        ("Bug bounty and abuse reporting. ", True),
        ("Operate a stable intake address whose underlying inbox can be rolled without changing "
         "the public-facing alias.",
         False),
    ])
    add_bullet(doc, [
        ("Privacy-preserving customer support. ", True),
        ("Give each ticket a scoped inbox; when the ticket closes, the inbox and all attachments "
         "are removed by the hourly scheduled() handler.",
         False),
    ])

    # ---- 5 Why ----
    add_section_heading(doc, "5. Why This Matters")
    add_bullet(doc, [
        ("Unlinkability by construction. ", True),
        ("The relay's two endpoint domains are independent on the wire. Correlation requires "
         "compromising the Worker itself, not either user's ISP.",
         False),
    ])
    add_bullet(doc, [
        ("Zero retention by default. ", True),
        ("The hourly scheduled() handler purges expired inboxes, R2 objects, KV sessions, and "
         "relay aliases in one pass. Retention is opt-in, not opt-out.",
         False),
    ])
    add_bullet(doc, [
        ("Native Cloudflare bindings end to end. ", True),
        ("MX, outbound, SQL, blob, KV, and real-time push are all platform primitives. There is "
         "no external SMTP relay, no external database, and no message queue to operate.",
         False),
    ])
    add_bullet(doc, [
        ("Authenticated and rate-limited at the edge. ", True),
        ("Bearer tokens for inbox access, password-authenticated admin sessions for domain "
         "registration, one-time WebSocket tickets with origin validation, and Turnstile on "
         "every creation surface.",
         False),
    ])
    add_bullet(doc, [
        ("Linear scale, no regional infrastructure. ", True),
        ("The Worker runs in every Cloudflare data center. Throughput grows with the network, "
         "not with provisioned capacity.",
         False),
    ])

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
